#! /usr/bin/env python3
from docx import Document

doc = Document()
doc.add_heading('Titolo', 0)
par = doc.add_paragraph('Un pò di testo normale, ')
par.add_run('e un pò di grassetto').bold = True

doc.add_heading('Altro titolo', 1)
doc.add_paragraph('Quotazione', 'Intense Quote')

par = doc.add_paragraph('E infine una tabella')
records = (
    (1, 'Roma'),
    (2, 'Barcellona'),
    (3, 'Parigi')
)
table = doc.add_table(1, 2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Id'
hdr_cells[1].text = 'Città'
for id, citta in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(id)
    row_cells[1].text = citta

doc.add_page_break()

doc.save('test.docx')
