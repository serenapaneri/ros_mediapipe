#! /usr/bin/env python3

import rospy
import time
from docx import Document

found = 0

def main():

    global found
    rospy.init_node('prova_file')
    found += 1

    start_time = time.time()
    doc = Document()
    doc.add_heading('ID: 000' + str(found), 0)

    doc.add_heading('Consciousness/Unconsciousness', 1)

    table1 = doc.add_table(1, 2)
    cells0 = table1.rows[0].cells
    cells0[0].text = 'Pose'
    cells0[1].text = 'Sit down'

    cells1 = table1.add_row().cells
    cells1[0].text = 'Motion'
    cells1[1].text = 'True'
    cells2 = table1.add_row().cells
    cells2[0].text = 'Eye-blinking'
    cells2[1].text = 'False'
    cells3 = table1.add_row().cells
    cells3[0].text = 'Response'
    cells3[1].text = 'True'


    doc.add_heading('Trauma/Non-trauma', 1)

    table2 = doc.add_table(1, 2)
    cells02 = table2.rows[0].cells
    cells02[0].text = 'Question 1'
    cells02[1].text = 'Clear answer'
    cells12 = table2.add_row().cells
    cells12[0].text = 'Question 2'
    cells12[1].text = 'Confused answer'
    cells22 = table2.add_row().cells
    cells22[0].text = 'Question 3a'
    cells22[1].text = 'Clear answer'
    cells32 = table2.add_row().cells
    cells32[0].text = 'Question 3b'
    cells32[1].text = 'True'
    
    doc.add_page_break()

    doc.save('/home/spot/ros_ws/src/spot_mediapipe/results/prova{}.docx'.format(found))

    stop_time = time.time()
    time_needed = stop_time - start_time

    print(time_needed)

    print(found)

    rospy.spin()

if __name__ == '__main__':
    main()
