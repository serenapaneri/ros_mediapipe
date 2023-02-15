#!/usr/bin/env python3

import rospy
import time
from docx import Document
from spot_mediapipe.srv import Trauma
from spot_mediapipe.srv import Blink
from spot_mediapipe.srv import Movement
from spot_mediapipe.srv import HumanPose
from spot_mediapipe.srv import Docu
from spot_mediapipe.srv import Found

# document service
docu_srv = None
# trauma client to recieve the answers
trauma_client = None 
# motion client to recieve information about the movement
motion_client = None 
# motion client to recieve information about the blinking
blink_client = None
# client to recieve information about the pose 
pose_client = None

start = False

##
# \brief Callback function of audio_srv
# \param req, AudioRequest
# \return start
#
# This callback function allows the client to start and stop the behavior
def doc_handle(req):
    global start
    if (req.docu == 'starting'):
        start = True
    elif (req.docu == 'stopping'):
        start = False
    return start

def main():

    global docu_srv, trauma_client, motion_client, blink_client, pose_client, start
    rospy.init_node('file')

    # document service
    docu_srv = rospy.Service('docum', Docu, doc_handle)
    # people client
    rospy.wait_for_service('foundpeople')
    people_client = rospy.ServiceProxy('foundpeople', Found)
    # motion client
    rospy.wait_for_service('movement')
    motion_client = rospy.ServiceProxy('movement', Movement)
    # pose client
    rospy.wait_for_service('human_pose')
    pose_client = rospy.ServiceProxy('human_pose', HumanPose)
    # blink client
    rospy.wait_for_service('blink')
    blink_client = rospy.ServiceProxy('blink', Blink)
    # trauma client
    trauma_client = rospy.ServiceProxy('trauma_questions', Trauma)

    rate = rospy.Rate(20)

    while not rospy.is_shutdown():

        if start == True:
            # number of people
            res_num = people_client()
            found = res_num.peoplefound
        
            doc = Document()
            doc.add_heading('ID: 000' + str(found), 0)

            # pose client
            res_pose = pose_client()
            pose = res_pose.pose_

            # trauma client
            res_trauma = trauma_client()
            q1 = res_trauma.question1
            q2 = res_trauma.question2
            q3a = res_trauma.question3a
            q3b = res_trauma.question3b

            # motion client
            res_motion = motion_client()
            motion = res_motion.mot

            # blink client
            res_blink = blink_client()
            blink = res_blink.blinking

            # creating consciousness/unconsciousness table
            doc.add_heading('Consciousness/Unconsciousness', 1)
            table_c = doc.add_table(1, 2)

            # pose
            cell_c1 = table_c.rows[0].cells
            cell_c1[0].text = 'Pose'
            if pose == 1:
                cell_c1[1].text = 'Stand up'
            elif pose == 2:
                cell_c1[1].text = 'Sit down'
            elif pose == 3:
                cell_c1[1].text = 'Lay down'
            elif pose == 0:
                cell_c1[1].text = 'Not detected'

            # motion
            cell_c2 = table_c.add_row().cells
            cell_c2[0].text = 'Motion'
            cell_c2[1].text = str(motion)

            # blink
            cell_c3 = table_c.add_row().cells
            cell_c3[0].text = 'Eye-blinking'
            cell_c3[1].text = str(blink)

            # talk
            cell_c4 = table_c.add_row().cells
            cell_c4[0].text = 'Response'
            if q1 == 1 or q1 == 2 or q2 == 1 or q2 == 2 or q3a == 1 or q3a == 2:
                cell_c4[1].text = 'True'
            else:
                cell_c4[1].text = 'False'


            doc.add_heading('Trauma/Non-Trauma', 1)
            table_t = doc.add_table(1, 2)

            # first question
            cell_t1 = table_t.rows[0].cells
            cell_t1[0].text = 'Question 1'
            if q1 == 1:
                cell_t1[1].text = 'Clear answer'
            elif q1 == 2:
                cell_t1[1].text = 'Confused answer'
            elif q1 == 3:
                cell_t1[1].text = 'No answer'

            # second question
            cell_t2 = table_t.add_row().cells
            cell_t2[0].text = 'Question 2'
            if q2 == 1:
                cell_t2[1].text = 'Clear answer'
            elif q2 == 2:
                cell_t2[1].text = 'Confused answer'
            elif q2 == 3:
                cell_t2[1].text = 'No answer'


            # third question
            cell_t3 = table_t.add_row().cells
            cell_t3[0].text = 'Question 3'
            if q3a == 1:
                cell_t3[1].text = 'Clear answer'
            elif q3a == 2:
                cell_t3[1].text = 'Confused answer'
            elif q3a == 3:
                cell_t3[1].text = 'No answer'


            # trauma question
            cell_t4 = table_t.add_row().cells
            cell_t4[0].text = 'Trauma'
            if q3b == True:
                cell_t4[1].text = 'True'
            elif q3b == False:
                cell_t4[1].text = 'False'


            doc.add_page_break()
            doc.save('/home/spot/ros_ws/src/spot_mediapipe/results/ID000{}.docx'.format(found))

            time.sleep(5)

            rate.sleep()

        elif start == False:
            rate.sleep()


if __name__ == '__main__':
    main()

