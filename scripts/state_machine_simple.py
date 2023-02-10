#! /usr/bin/env python3

# This script is simplified since the exploration part is not done with the rrt navigation but it
# is only done with a Twist and Odometry message. And the assuntion here is single person in a room.

import rospy
import random
import smach
import smach_ros
import time
from docx import Document
from geometry_msgs.msg import Twist
from geometry_msgs.msg import Pose
from geometry_msgs.msg import Point
from nav_msgs.msg import Odometry
from spot_mediapipe.srv import Audio
from spot_mediapipe.srv import Trauma
from spot_mediapipe.srv import Blink
from spot_mediapipe.srv import Movement
from spot_mediapipe.srv import HumanPose

# audio client to start the trauma behavior
audio_client = None 
# trauma client to recieve the answers
trauma_client = None 
# motion client to recieve information about the movement
motion_client = None 
# motion client to recieve information about the blinking
blink_client = None
# client to recieve information about the pose 
pose_client = None
# detection client
detection_client = None

# velocity publisher
pub_cmd = None
# body pose publisher
pub_body = None

# robot state variables
actual_position = Point()
actual_position.x = 0 
actual_position.y = 0 
actual_yaw = 0

# list of landmarks collected, len(landmarks) is always 33
landmarks = []

# poses in which the person can be found 
pose = 0 
human_pose = 0
detection = False

# people found
found_people = 0 

stop = False

##### PARAMETERS #####

# velocity parameters
lin_vel = 0.5 # you should set this parameter
ang_vel = 0.5 # you should set this parameter

# body pose parameters
#### DEVI SETTARE NUOVI PARAMETRI
default = 0.0
look_up = 0.1415
look_down = 0.08 

# proxemics parameters
distance_standup = 0 ## DA SETTARE
distance_sitdown = 0 ## DA SETTARE
distance_laydown = 0 ## DA SETTARE

# point where the person is
goal_position_x = 0
goal_position_y = 0

start_time = 0


##
# This is the callback function of the subscriber to the topic odom and it is used to know the actual position
# and orientation of the robot in the space.
def odom_callback(msg):
    global actual_position, actual_yaw
    
    # actual position
    actual_position = msg.pose.pose.position
    # actual yaw
    quaternion = (
        msg.pose.pose.orientation.x,
        msg.pose.pose.orientation.y,
        msg.pose.pose.orientation.z,
        msg.pose.pose.orientation.w)
    euler = transformations.euler_from_quaternion(quaternion)
    actual_yaw = euler[2]



##
# Function to go on a straight line for a certain distance
def go_forward(distance):

    global pub_cmd, lin_vel
    twist_msg = Twist()
    twist_msg.linear.x = lin_vel
    twist_msg.angular.z = 0

    start_time = rospy.Time.now().to_sec()
    current_distance = 0

    while (current_distance < distance):
        pub_cmd.publish(twist_msg)
        time = rospy.Time.now().to_sec()
        current_distance = lin_vel*(time - start_time)

    twist_msg.linear.x = 0
    pub_cmd.publish(twist_msg)
    print('The robot is moving forward')


##
# Function to make the robot rotate counterclockwise
### CREDO SIA IN GRADI
def yaw(angle):

    global pub_cmd, ang_vel
    relative_angle = angle*(2*math.pi)/360

    twist_msg = Twist()
    twist_msg.linear.x = 0
    twist_msg.angular.z = ang_vel

    start_time = rospy.Time.now().to_sec()
    current_angle = 0

    while (current_angle < relative_angle):
        pub_cmd.publish(twist_msg)
        time = rospy.Time.now().to_sec()
        current_angle = ang_vel*(time - start_time)
    twist_msg.angular.z = 0
    pub_cmd.publish(twist_msg)
    print('The robot is rotating')


def pose_body(pose):

    global pub_body
    pose_msg = Pose()
    pose_msg.position.x = 0.0
    pose_msg.position.y = 0.0
    pose_msg.position.z = 0.0
    pose_msg.orientation.x = 0.0  
    pose_msg.orientation.y = pose
    pose_msg.orientation.z = 0.0
    pose_msg.orientation.w = 1.0

    pub_body.publish(pose_msg)
    print('The robot is looking up')


##
# Function used to stop the robot's behavior
def stop():

    global pub_cmd
    twist_msg = Twist()
    twist_msg.linear.x = 0
    twist_msg.angular.z = 0
    pub_cmd.publish(twist_msg)
    print('The robot is stopping')



##
# In this class the robot should simulate the exploration part, so it navigates randomly in the 
# environment, while building a map and searching for people to rescue.
# When someone is detected then the state change and the robot computes the pose of the person
# to better approach it and then starts all the computations.
class Motion(smach.State):

    def __init__(self):
        smach.State.__init__(self, 
                             outcomes=['motion','compute_pose'])
        
    def execute(self, userdata):
     
       global detection_client, pub_cmd, landmarks, detection, found_people

       # retrieving, if there are any the mediapipe markers
       res_detection = detection_client()
       detection = res_detection.detected

       # if a person is detected
       if detection:
           found_people += 1
           # creating the custom file for the person just found
           doc = Document()
           doc.add_heading('ID: 000' + str(found_people), 0)
           start_time = time.time()  
           return 'compute_pose'

       # if there is nobody
       else:
           # here wuold go the exploration part, it is simplify as
           ### DOES NOT WORK BUT THE IDEA IS THAT ONE
           go_forward(3)
           stop()
           yaw(45)
           stop()
           
           return 'motion'

           

##
# In this class the robot computes the pose of the person detected in order to better approaches it.
class ComputePose(smach.State):

    def __init__(self):
        smach.State.__init__(self, 
                             outcomes=['go_to'])
        
    def execute(self, userdata):

        global pose_client, pose

        res_pose = pose_client()
        pose = res_pose.pose_

        # stand up
        if pose == 1:
            print('The person is standing up')
        # sit down
        elif pose == 2:
            print('The person is sitting down')
        # lay down
        elif pose == 3:
            print('The person is laying down')
   
        return 'go_to'


##
# In this class the robot approaches the person, knowing in advance its pose, to better compute the
# position of the robot where all the computations are done.
class GoTo(smach.State):

    def __init__(self):
        smach.State.__init__(self, 
                             outcomes=['adjust'])
        
    def execute(self, userdata):

        global pose, distance_standup, distance_sitdown, distance_laydown, actual_position, goal_position_x, goal_position_y
       
        # implement move base

        # if the person is standing up
        if pose == 1:
            while ((actual_position.x - goal_position_x)*(actual_position.x - goal_position_x) + (actual_position.y - goal_position_y)*(actual_position.y - goal_position_y)) > distance_standup:
                time.sleep(0.1)    

        # if the person is sitting down
        elif pose == 2:
            while ((actual_position.x - goal_position_x)*(actual_position.x - goal_position_x) + (actual_position.y - goal_position_y)*(actual_position.y - goal_position_y)) > distance_sitdown:
                time.sleep(0.1) 

        # if the person is laying down
        elif pose == 3:
            while ((actual_position.x - goal_position_x)*(actual_position.x - goal_position_x) + (actual_position.y - goal_position_y)*(actual_position.y - goal_position_y)) > distance_laydown:
                time.sleep(0.1) 
    
        return 'adjust'

##
# In this class the robot adjust its own position to have an optimal view of the person in order to
# starts all the analisys to be done. In the adjustment the robot should also assume a specific pose
# on the basis of the pose of the person and then stop every kind of motion. 
class Adjust(smach.State):

    def __init__(self):
        smach.State.__init__(self, 
                             outcomes=['analisys'])
        
    def execute(self, userdata):
    
        global default, look_up, look_down, pub_body

                 

        return 'analisys'


##
# In this class the robot verify both the consciousness of the people and the trauma computation. These
# two kinds of analisy are done simultaneosly in order to reduce the rescue time, which is crucial for
# these kind of applications. Finally the robot should   
class Analisys(smach.State):

    def __init__(self):
        smach.State.__init__(self, 
                             outcomes=['everything_checked'])
        
    def execute(self, userdata):

        global audio_client, trauma_client, motion_client, blink_client, pose, stop, start_time

        audio_client('start')

        # trauma
        res_trauma = trauma_client()
        stop = res_trauma.ok

        ### CONTROLLARE PERCHE MI PUZZA
        while stop == False:
            q1 = res_trauma.question1
            q2 = res_trauma.question2
            q3a = res_trauma.question3a
            q3b = res_trauma.question3b
            # stop = res_trauma.ok
            time.sleep(0.1)

        # verifica se prende le domande
        audio_client('stop')

        # motion client
        res_motion = motion_client()
        motion = res_motion.mot
        print('Motion: {}'.format(motion))

        # blink client
        res_blink = blink_client()
        blink = res_blink.blinking
        print('Eye-blinking: {}'.format(blink))

        # creating consciousness/unconsciousness table
        doc.add_heading('Consciousness/Unconsciousness', 1)
        table_c = doc.add_table(1, 2)

        # pose
        cell_c1 = table_c.rows[0].cells
        cell_c1[0].text = 'Pose'
        if pose == 1:
            cell_c1[1].text = 'Stand up'
        elif pose == 2:
            cell_c1[1].text = 'Sit down'
        elif pose == 3:
            cell_c1[1].text = 'Lay down'

        # motion
        cell_c2 = table_c.add_row().cells
        cell_c2[0].text = 'Motion'
        cell_c2[1].text = str(motion)

        # blink
        cell_c3 = table_c.add_row().cells
        cell_c3[0].text = 'Eye-blinking'
        cell_c3[1].text = str(blink)

        # talk
        cell_c4 = table_c.add_row().cells
        cell_c4[0].text = 'Response'
        if q1 == 1 or q1 == 2 or q2 == 1 or q2 == 2 or q3a == 1 or q3a == 2:
            print('Response: True')
            cell_c4[1].text = 'True'
        else:
            print('Response: False')
            cell_c4[1].text = 'False'


        doc.add_heading('Trauma/Non-Trauma', 1)
        table_t = doc.add_table(1, 2)

        # first question
        cell_t1 = table_t.rows[0].cells
        cell_t1[0].text = 'Question 1'
        if q1 == 1:
            print('Question1: Clear answer')
            cell_t1[1].text = 'Clear answer'
        elif q1 == 2:
            print('Question1: Confused answer')
            cell_t1[1].text = 'Confused answer'
        elif q1 == 3:
            print('Question1: No answer')
            cell_t1[1].text = 'No answer'

        # second question
        cell_t2 = table_t.add_row().cells
        cell_t2[0].text = 'Question 2'
        if q2 == 1:
            print('Question2: Clear answer')
            cell_t2[1].text = 'Clear answer'
        elif q2 == 2:
            print('Question2: Confused answer')
            cell_t2[1].text = 'Confused answer'
        elif q2 == 3:
            print('Question2: No answer')
            cell_t2[1].text = 'No answer'


        # third question
        cell_t3 = table_t.add_row().cells
        cell_t3[0].text = 'Question 3'
        if q3a == 1:
            print('Question3: Clear answer')
            cell_t3[1].text = 'Clear answer'
        elif q3a == 2:
            print('Question3: Confused answer')
            cell_t3[1].text = 'Confused answer'
        elif q3a == 3:
            print('Question3: No answer')
            cell_t3[1].text = 'No answer'


        # trauma question
        cell_t4 = table_t.add_row().cells
        cell_t4[0].text = 'Trauma'
        if q3b == True:
            print('Trauma: True')
            cell_t4[1].text = 'True'
        elif q3b == False:
            print('Trauma: False')
            cell_t4[1].text = 'False'

        stop_time = time.time()
        time_needed = stop_time - start_time

        print('Rescue time: {}'.format(time_needed))
        par = doc.add_paragraph('Rescue time: {}'.format(time_needed)).bold = True

        doc.add_page_break()
        doc.save('/home/spot/ros_ws/src/spot_mediapipe/results/ID000{}.docx'.format(found_people))

        return 'everything_checked' 
        

def main():

    global audio_client, trauma_client, motion_client, pose_client, detection_client, blink_client, pub_cmd, pub_body
    rospy.init_node('state_machine_simple')
    sm = smach.StateMachine(outcomes=['everything_checked'])

    # detection client
    detection_client = rospy.ServiceProxy('detection', Detection)
    # audio client
    audio_client = rospy.ServiceProxy('trauma_audio', Audio)
    # trauma client
    trauma_client = rospy.ServiceProxy('trauma_questions', Trauma)
    # motion client
    motion_client = rospy.ServiceProxy('movement', Movement)
    # pose client
    pose_client = rospy.ServiceProxy('human_pose', HumanPose)
    # blink client
    blink_client = rospy.ServiceProxy('blink', Blink)
    # cmd_vel publisher
    pub_cmd = rospy.Publisher('/spot/cmd_vel', Twist, queue_size = 1)
    # body_pose publisher
    pub_body = rospy.Publisher('/spot/body_pose', Pose, queue_size = 1)

    # CLIENT FOR MOVEBASE


    with sm:
        smach.StateMachine.add('Motion', Motion(), 
                               transitions={'compute_pose':'ComputePose',
                                            'motion':'Motion'})
        smach.StateMachine.add('ComputePose', ComputePose(), 
                               transitions={'go_to':'GoTo'})
        smach.StateMachine.add('GoTo', GoTo(), 
                               transitions={'adjust':'Adjust'})
        smach.StateMachine.add('Adjust', GoTo(), 
                               transitions={'analisys':'Analisys'})
        smach.StateMachine.add('Analisys', Analisys(), 
                               transitions={'everything_checked':'everything_checked'})


    # Create and start the introspection server for visualization
    sis = smach_ros.IntrospectionServer('state_machine', sm, '/SM_ROOT')
    sis.start()

    # Execute the state machine
    outcome = sm.execute()

    # Wait for ctrl-c to stop the application
    rospy.spin()
    sis.stop()
    
if __name__ == '__main__':
    main()


